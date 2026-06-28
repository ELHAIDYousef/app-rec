"""F13 – Génération automatique du cahier des charges + export PDF/Word"""
from io import BytesIO
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.core.security import require_role
from app.core.groq_helper import appeler_groq_json
from app.models.stage import CandidatureStage, CahierDesCharges, ProfilAnalyse

router = APIRouter(prefix="/api/stage/cahier", tags=["Stage – F13 Cahier des charges"])


# ── Sérialisation ─────────────────────────────────────────────

def _ser(c: CahierDesCharges) -> dict:
    return {
        "id":             c.id,
        "candidature_id": c.candidature_id,
        "sujet_id":       c.sujet_id,
        "sujet_titre":    c.sujet.titre if c.sujet else None,
        "contenu":        c.contenu or {},
        "statut":         c.statut,
        "cree_le":        c.cree_le.isoformat() if c.cree_le else None,
    }


# ── Génération PDF ────────────────────────────────────────────

def _build_pdf(cahier: CahierDesCharges) -> BytesIO:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    BLUE       = colors.HexColor("#1D4ED8")
    GRAY       = colors.HexColor("#6B7280")
    LIGHT_GRAY = colors.HexColor("#F9FAFB")
    BORDER     = colors.HexColor("#E5E7EB")
    TEXT       = colors.HexColor("#374151")

    s_title  = ParagraphStyle("title",  fontSize=20, fontName="Helvetica-Bold",
                               alignment=TA_CENTER, textColor=BLUE, spaceAfter=4)
    s_sub    = ParagraphStyle("sub",    fontSize=14, fontName="Helvetica-Bold",
                               alignment=TA_CENTER, textColor=TEXT, spaceAfter=4)
    s_meta   = ParagraphStyle("meta",   fontSize=9,  fontName="Helvetica",
                               alignment=TA_CENTER, textColor=GRAY, spaceAfter=2)
    s_h2     = ParagraphStyle("h2",     fontSize=11, fontName="Helvetica-Bold",
                               textColor=colors.white, spaceBefore=10, spaceAfter=0,
                               backColor=BLUE, leftIndent=-0.2*cm, rightIndent=-0.2*cm,
                               borderPadding=(4, 8, 4, 8))
    s_body   = ParagraphStyle("body",   fontSize=9,  fontName="Helvetica",
                               textColor=TEXT, leading=14, spaceAfter=2)
    s_bullet = ParagraphStyle("bullet", fontSize=9,  fontName="Helvetica",
                               textColor=TEXT, leading=14, leftIndent=12,
                               bulletIndent=4, spaceAfter=2)

    c = cahier.contenu or {}
    story = []

    # En-tête
    story.append(Paragraph("CAHIER DES CHARGES", s_title))
    if cahier.sujet:
        story.append(Paragraph(cahier.sujet.titre, s_sub))
    date_str = cahier.cree_le.strftime("%d/%m/%Y") if cahier.cree_le else ""
    story.append(Paragraph(f"Genere le {date_str}", s_meta))
    if c.get("duree_stage"):
        story.append(Paragraph(f"Duree : {c['duree_stage']}", s_meta))
    story.append(Spacer(1, 0.4 * cm))

    # Ligne de séparation
    story.append(Table([[""]], colWidths=[17 * cm], rowHeights=[2],
                       style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), BLUE)])))
    story.append(Spacer(1, 0.4 * cm))

    tbl_style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BLUE),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("GRID",       (0, 0), (-1, -1), 0.5, BORDER),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_GRAY]),
        ("PADDING",    (0, 0), (-1, -1), 6),
        ("VALIGN",     (0, 0), (-1, -1), "TOP"),
    ])

    def add_section(num, title, content):
        if not content:
            return
        story.append(Paragraph(f"{num}. {title}", s_h2))
        story.append(Spacer(1, 0.2 * cm))
        if isinstance(content, list):
            for item in content:
                story.append(Paragraph(f"- {str(item)}", s_bullet))
        elif isinstance(content, str):
            story.append(Paragraph(content, s_body))
        story.append(Spacer(1, 0.3 * cm))

    add_section(1, "Presentation du Projet", c.get("presentation"))
    add_section(2, "Contexte", c.get("contexte"))
    add_section(3, "Objectifs", c.get("objectifs"))
    add_section(4, "Perimetre Fonctionnel", c.get("perimetre_fonctionnel"))

    techs = c.get("technologies_recommandees")
    if techs:
        story.append(Paragraph("5. Technologies Recommandees", s_h2))
        story.append(Spacer(1, 0.2 * cm))
        data = [["Domaine", "Technologies"]]
        for k, v in techs.items():
            vals = ", ".join(v) if isinstance(v, list) else str(v)
            data.append([k.replace("_", " ").capitalize(), vals])
        story.append(Table(data, colWidths=[4.5 * cm, 12.5 * cm], style=tbl_style))
        story.append(Spacer(1, 0.3 * cm))

    add_section(6, "Livrables", c.get("livrables"))

    planning = c.get("planning")
    if planning:
        story.append(Paragraph("7. Planning", s_h2))
        story.append(Spacer(1, 0.2 * cm))
        data = [["Semaine", "Description"]]
        for p in planning:
            data.append([f"Sem. {p.get('semaine', '')}", p.get("description", "")])
        story.append(Table(data, colWidths=[3 * cm, 14 * cm], style=tbl_style))
        story.append(Spacer(1, 0.3 * cm))

    add_section(8, "Criteres d'Evaluation", c.get("criteres_evaluation"))
    add_section(9, "Methodologie", c.get("methodologie"))

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            rightMargin=2 * cm, leftMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    doc.build(story)
    buf.seek(0)
    return buf


# ── Génération Word ───────────────────────────────────────────

def _build_word(cahier: CahierDesCharges) -> BytesIO:
    from docx import Document as WordDoc
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLUE = RGBColor(0x1D, 0x4E, 0xD8)
    GRAY = RGBColor(0x6B, 0x72, 0x80)

    def set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    doc = WordDoc()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Titre
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CAHIER DES CHARGES")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE

    # Sous-titre
    if cahier.sujet:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cahier.sujet.titre)
        run.bold = True
        run.font.size = Pt(14)

    c = cahier.contenu or {}

    # Meta
    date_str = cahier.cree_le.strftime("%d/%m/%Y") if cahier.cree_le else ""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Généré le {date_str}")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    if c.get("duree_stage"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Durée : {c['duree_stage']}")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

    doc.add_paragraph()

    def add_heading(text: str):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = BLUE

    def add_content(content):
        if isinstance(content, list):
            for item in content:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(str(item))
                run.font.size = Pt(10)
        elif isinstance(content, str):
            p = doc.add_paragraph(content)
            for r in p.runs:
                r.font.size = Pt(10)

    def add_section(num: int, title: str, content):
        if not content:
            return
        add_heading(f"{num}. {title}")
        add_content(content)

    def add_table(headers: list, rows: list):
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            set_cell_bg(hdr_cells[i], "1D4ED8")
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for row_data in rows:
            row_cells = tbl.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = val
                for para in row_cells[i].paragraphs:
                    for r in para.runs:
                        r.font.size = Pt(10)

    add_section(1, "Présentation du Projet", c.get("presentation"))
    add_section(2, "Contexte", c.get("contexte"))
    add_section(3, "Objectifs", c.get("objectifs"))
    add_section(4, "Périmètre Fonctionnel", c.get("perimetre_fonctionnel"))

    techs = c.get("technologies_recommandees")
    if techs:
        add_heading("5. Technologies Recommandées")
        rows = []
        for k, v in techs.items():
            vals = ", ".join(v) if isinstance(v, list) else str(v)
            rows.append([k.replace("_", " ").capitalize(), vals])
        add_table(["Domaine", "Technologies"], rows)
        doc.add_paragraph()

    add_section(6, "Livrables", c.get("livrables"))

    planning = c.get("planning")
    if planning:
        add_heading("7. Planning")
        rows = [[f"Sem. {p.get('semaine', '')}", p.get("description", "")] for p in planning]
        add_table(["Semaine", "Description"], rows)
        doc.add_paragraph()

    add_section(8, "Critères d'Évaluation", c.get("criteres_evaluation"))
    add_section(9, "Méthodologie", c.get("methodologie"))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Helpers nom de fichier ────────────────────────────────────

def _nom(cahier: CahierDesCharges) -> str:
    titre = (cahier.sujet.titre if cahier.sujet else "stage") or "stage"
    return titre.replace(" ", "_")[:40]


# ── Endpoints CRUD ────────────────────────────────────────────

@router.post("/generer/{cid}")
def generer(
    cid: int,
    db: Session = Depends(get_db),
    user=Depends(require_role("admin", "rh")),
):
    cand = db.get(CandidatureStage, cid)
    if not cand: raise HTTPException(404, "Candidature introuvable")
    if not cand.sujet: raise HTTPException(400, "Aucun sujet affecté à cette candidature")

    profil = db.query(ProfilAnalyse).filter(ProfilAnalyse.candidature_id == cid).first()
    s = cand.sujet

    prompt = f"""Tu es expert en gestion de projets informatiques. Génère un cahier des charges complet pour ce stage.

SUJET : {s.titre}
DESCRIPTION : {s.description[:800]}
TECHNOLOGIES : {s.technologies or []}
NIVEAU REQUIS : {s.niveau_requis or 'Non précisé'}
ENCADRANT : {s.encadrant or 'Non précisé'}
PROFIL STAGIAIRE : Niveau {profil.niveau_technique if profil else 'intermédiaire'}, compétences: {profil.competences[:5] if profil and profil.competences else []}

Retourne ce JSON complet :
{{
  "presentation": "Présentation du projet en 3-4 phrases",
  "contexte": "Contexte et enjeux du projet",
  "objectifs": ["objectif 1", "objectif 2", "objectif 3"],
  "perimetre_fonctionnel": ["fonctionnalité 1", "fonctionnalité 2", "fonctionnalité 3", "fonctionnalité 4"],
  "technologies_recommandees": {{"frontend": ["tech1"], "backend": ["tech2"], "base_donnees": ["tech3"], "outils": ["outil1"]}},
  "livrables": ["livrable 1", "livrable 2", "livrable 3"],
  "planning": [
    {{"semaine": "1-2", "description": "Phase d'analyse et de cadrage"}},
    {{"semaine": "3-4", "description": "Conception et architecture"}},
    {{"semaine": "5-8", "description": "Développement des fonctionnalités principales"}},
    {{"semaine": "9-10", "description": "Tests et corrections"}},
    {{"semaine": "11-12", "description": "Déploiement et documentation"}}
  ],
  "criteres_evaluation": ["critère 1", "critère 2", "critère 3"],
  "methodologie": "Description de la méthodologie (Scrum, sprints de 2 semaines, etc.)",
  "duree_stage": "3 mois"
}}"""

    contenu = appeler_groq_json([
        {"role": "system", "content": "Tu es expert en cahiers des charges de projets informatiques. Réponds uniquement en JSON valide et complet."},
        {"role": "user", "content": prompt},
    ], max_tokens=2000)

    old = db.query(CahierDesCharges).filter(CahierDesCharges.candidature_id == cid).first()
    if old: db.delete(old)

    cahier = CahierDesCharges(
        stagiaire_id   = cand.stagiaire_id,
        candidature_id = cid,
        sujet_id       = cand.sujet_id,
        contenu        = contenu,
        statut         = "genere",
    )
    db.add(cahier); db.commit(); db.refresh(cahier)
    return _ser(cahier)


@router.get("/mon-cahier")
def mon_cahier(db: Session = Depends(get_db), user=Depends(require_role("stagiaire"))):
    cand = db.query(CandidatureStage).filter(
        CandidatureStage.stagiaire_id == user.id,
        CandidatureStage.statut == "acceptee",
    ).first()
    if not cand: raise HTTPException(404, "Aucune candidature acceptée")

    cahier = db.query(CahierDesCharges).filter(CahierDesCharges.candidature_id == cand.id).first()
    if not cahier: raise HTTPException(404, "Cahier des charges non encore généré")
    return _ser(cahier)


# ── Export stagiaire ──────────────────────────────────────────

@router.get("/mon-cahier/pdf")
def mon_cahier_pdf(db: Session = Depends(get_db), user=Depends(require_role("stagiaire"))):
    cand = db.query(CandidatureStage).filter(
        CandidatureStage.stagiaire_id == user.id,
        CandidatureStage.statut == "acceptee",
    ).first()
    if not cand: raise HTTPException(404, "Aucune candidature acceptée")
    cahier = db.query(CahierDesCharges).filter(CahierDesCharges.candidature_id == cand.id).first()
    if not cahier: raise HTTPException(404, "Cahier des charges non encore généré")

    buf = _build_pdf(cahier)
    return StreamingResponse(
        buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="cahier_{_nom(cahier)}.pdf"'},
    )


@router.get("/mon-cahier/word")
def mon_cahier_word(db: Session = Depends(get_db), user=Depends(require_role("stagiaire"))):
    cand = db.query(CandidatureStage).filter(
        CandidatureStage.stagiaire_id == user.id,
        CandidatureStage.statut == "acceptee",
    ).first()
    if not cand: raise HTTPException(404, "Aucune candidature acceptée")
    cahier = db.query(CahierDesCharges).filter(CahierDesCharges.candidature_id == cand.id).first()
    if not cahier: raise HTTPException(404, "Cahier des charges non encore généré")

    buf = _build_word(cahier)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="cahier_{_nom(cahier)}.docx"'},
    )


# ── Export admin/RH ───────────────────────────────────────────

@router.get("/{id}/pdf")
def cahier_pdf(id: int, db: Session = Depends(get_db), user=Depends(require_role("admin", "rh"))):
    cahier = db.get(CahierDesCharges, id)
    if not cahier: raise HTTPException(404)
    buf = _build_pdf(cahier)
    return StreamingResponse(
        buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="cahier_{_nom(cahier)}.pdf"'},
    )


@router.get("/{id}/word")
def cahier_word(id: int, db: Session = Depends(get_db), user=Depends(require_role("admin", "rh"))):
    cahier = db.get(CahierDesCharges, id)
    if not cahier: raise HTTPException(404)
    buf = _build_word(cahier)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="cahier_{_nom(cahier)}.docx"'},
    )


# ── Autres endpoints ──────────────────────────────────────────

@router.get("/{id}")
def obtenir(id: int, db: Session = Depends(get_db), user=Depends(require_role("admin", "rh"))):
    c = db.get(CahierDesCharges, id)
    if not c: raise HTTPException(404)
    return _ser(c)


@router.patch("/{id}/valider")
def valider(id: int, db: Session = Depends(get_db), user=Depends(require_role("admin", "rh"))):
    c = db.get(CahierDesCharges, id)
    if not c: raise HTTPException(404)
    c.statut = "valide"
    db.commit(); db.refresh(c)
    return _ser(c)
